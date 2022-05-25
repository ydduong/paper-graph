# -*- coding: utf-8 -*-

from unit import Args, baidu_trans
import openpyxl

# 导入库 python-docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches
from docx.oxml.ns import qn

# 分词
from nltk import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer


class Write:
    def __init__(self, args: Args):
        self._check_func = args.check_is_keyword_in_strings
        self._excel_file = args.excel_file
        self._markdown = args.markdown
        self._docx = args.docx
        self._words_list = args.words_list
        self._log = args.log
        self._is_zh = args.is_zh

        self._sqlite = args.sqlite
        self._result = self._sqlite.select_all_from_paper()
        # self._header = self._sqlite.header

        # 索引
        self.header = {'title': 0, 'author': 1, 'year': 2, 'journal': 3, 'citations': 4, 'year_citations': 5,
                       'connected_papers_url': 6, 'semantic_scholar_url': 7, 'publisher_page_url': 8, 'abstract': 9,
                       'title_zh': 10, 'abstract_zh': 11}

        # 对结果按引用量排序
        self._result = list(self._result)
        self._result.sort(key=lambda x: float(x[5]), reverse=True)

        count = 0
        for item in self._result:
            if self._check_func(item[self.header.get('title')]):
                print(item)
                count += 1
            if count > 4:
                break

        pass

    def to_excel(self):
        self._log.append('Starting write to excel.')
        target_workbook = openpyxl.Workbook(write_only=True)
        target_sheet = target_workbook.create_sheet('sheet0')

        # 表头
        header = ['title', 'author', 'year', 'journal', 'citations', 'year_citations', 'connected_papers_url',
                  'semantic_scholar_url', 'publisher_page_url', 'abstract']
        if self._is_zh:
            header = ['title', 'author', 'year', 'journal', 'citations', 'year_citations', 'connected_papers_url',
                      'semantic_scholar_url', 'publisher_page_url', 'abstract', 'title_zh', 'abstract_zh']

        target_sheet.append(header)

        # 写入
        if self._is_zh:
            for item in self._result:
                item = list(item)
                if len(item[self.header.get('title_zh')]) == 0:
                    item[self.header.get('title_zh')] = baidu_trans(item[self.header.get('title')])

                if len(item[self.header.get('abstract_zh')]) == 0:
                    item[self.header.get('abstract_zh')] = baidu_trans(item[self.header.get('abstract')])

                if self._check_func(item[self.header.get('title')]):
                    target_sheet.append(item)

                # 翻译之后 需要 写入数据库
        else:
            for item in self._result:
                if self._check_func(item[self.header.get('title')]):
                    item = list(item)
                    item = item[:-2]
                    target_sheet.append(item)

        target_workbook.save(self._excel_file)
        pass

    def to_markdown(self):
        self._log.append('Starting write to markdown.')

        # 写入
        sign = 0
        with open(self._markdown, 'w', encoding='utf-8') as w:
            for item in self._result:
                if not self._check_func(item[self.header.get('title')]):
                    continue
                sign += 1

                if self._is_zh:  # 标题、链接、引用，年均引用，作者、年份，期刊、中文标题、中文摘要
#                     strs = f'''### {sign}.{item[self.header.get('title')]}
# {item[self.header.get('citations')]}, {item[self.header.get('year_citations')]}, {item[self.header.get('author')]}
# {item[self.header.get('year')]}, {item[self.header.get('journal')]}
# {item[self.header.get('title_zh')]}
# {item[self.header.get('abstract_zh')]}
#
# '''
                    strs = f'''### {sign}.{item[self.header.get('title')]}
- {item[self.header.get('publisher_page_url')]}
- {item[self.header.get('citations')]}, {item[self.header.get('year_citations')]}, {item[self.header.get('author')]}
- {item[self.header.get('year')]}, {item[self.header.get('journal')]}
- {item[self.header.get('title_zh')]}
- {item[self.header.get('abstract_zh')]}

'''
                else:  # 标题、链接、引用，年均引用，作者、年份，期刊、摘要
                    strs = f'''### {sign}.{item[self.header.get('title')]}
- {item[self.header.get('publisher_page_url')]}
- {item[self.header.get('citations')]}, {item[self.header.get('year_citations')]}, {item[self.header.get('author')]}
- {item[self.header.get('year')]}, {item[self.header.get('journal')]}
- {item[self.header.get('abstract')]}

'''
                w.write(strs)
        pass

    def to_word(self):
        # https://blog.csdn.net/auspark/article/details/106634417
        self._log.append('Starting write to docx.')

        # 新建空白文档
        doc = Document()

        # 正文样式
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'  # 必须先设置font.name
        style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        sign = 0
        for item in self._result:
            if not self._check_func(item[self.header.get('title')]):
                continue
            sign += 1

            if self._is_zh:  # 标题、链接、引用，年均引用，作者、年份，期刊、中文标题、中文摘要
                print('Warning: please check write to word code for zh.')
                pass
            else:  # 标题、链接、引用，年均引用，作者、年份，期刊、摘要
                title = doc.add_heading(level=1)
                title.paragraph_format.space_before = Pt(6)  # 设置段前 0 磅
                title.paragraph_format.space_after = Pt(6)  # 设置段后 0 磅
                title.paragraph_format.line_spacing = 1.2  # 设置行间距为 1.5
                run = title.add_run(f"{sign}.{item[self.header.get('title')]}")
                run.font.size = Pt(14)  # 设置字体大小，小四对应值为12
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = 'Times New Roman'  # 设置字体类型属性
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # font = u'黑体'

                # 增加无序列表
                doc.add_paragraph(item[self.header.get('publisher_page_url')], style='List Bullet')
                doc.add_paragraph(f"{item[self.header.get('citations')]}, {item[self.header.get('year_citations')]}, {item[self.header.get('author')]}", style='List Bullet')
                doc.add_paragraph(f"{item[self.header.get('year')]}, {item[self.header.get('journal')]}", style='List Bullet')
                paragraph = doc.add_paragraph({item[self.header.get('abstract')]}, style='List Bullet')
                paragraph.alignment = 3
                doc.add_paragraph()

        # 保存文件
        doc.save(self._docx)

    def to_title_word_list(self):
        # 记录词频
        stem = {}

        for item in self._result:
            if not self._check_func(item[self.header.get('title')]):
                continue

            # 以空格形式实现分词
            words = word_tokenize(item[self.header.get('title')])

            # 除去符号
            punctuation = [',', '.', ':', ';', '?', '(', ')', '[', ']', '&', '!', '*', '@', '#', '$', '%']  # 定义标点符号列表
            words = [word for word in words if word not in punctuation]  # 去除标点符号

            # 除去停用词
            stops = set(stopwords.words("english"))
            words = [word for word in words if word not in stops]
            # print(words)

            # 提取词干
            for word in words:
                word = PorterStemmer().stem(word)
                stem[word] = stem.get(word, 0) + 1  # 词干提取

        # 排序 [('heard', 3), ('i', 2)]
        res = sorted(stem.items(), key=lambda x: x[1], reverse=True)

        # 写入文件
        with open(self._words_list, 'w', encoding='utf-8') as w:
            for item in res:
                strs = f'{item[0]}, {item[1]}\n'
                w.write(strs)

    def run(self):
        self.to_title_word_list()
        self.to_excel()
        self.to_markdown()
        self.to_word()


if __name__ == '__main__':
    _args = Args()
    _w = Write(_args)
    _w.to_excel()
    pass
